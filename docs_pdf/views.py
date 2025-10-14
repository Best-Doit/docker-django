import os
import subprocess
import mimetypes
import cv2
import numpy as np
from django.shortcuts import render
from django.http import HttpResponse, HttpResponseNotFound
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from django.views.decorators.csrf import csrf_protect
from django.views.decorators.http import require_http_methods
from pdf2docx import Converter
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader

def preprocess_image_for_ocr(image_path):
    """
    Preprocesa una imagen para mejorar la precisión del OCR
    Usa PIL para compatibilidad con Railway
    """
    try:
        # Intentar usar OpenCV si está disponible
        try:
            import cv2
            import numpy as np
            
            # Leer imagen con OpenCV
            img = cv2.imread(image_path)
            if img is None:
                raise Exception("No se pudo cargar la imagen")
            
            # Convertir a escala de grises
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Reducir ruido
            denoised = cv2.medianBlur(gray, 3)
            
            # Aplicar filtro bilateral para reducir ruido manteniendo bordes
            filtered = cv2.bilateralFilter(denoised, 9, 75, 75)
            
            # Mejorar contraste usando CLAHE
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(filtered)
            
            # Aplicar umbralización adaptativa
            thresh = cv2.adaptiveThreshold(enhanced, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            # Operaciones morfológicas para limpiar la imagen
            kernel = np.ones((1,1), np.uint8)
            cleaned = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            # Guardar imagen preprocesada temporalmente
            temp_path = image_path.replace('.', '_processed.')
            cv2.imwrite(temp_path, cleaned)
            
            return temp_path
            
        except ImportError:
            # Fallback a PIL si OpenCV no está disponible
            print("OpenCV no disponible, usando PIL para preprocesamiento básico")
            
            # Usar PIL para preprocesamiento básico
            image = Image.open(image_path)
            
            # Convertir a escala de grises
            if image.mode != 'L':
                image = image.convert('L')
            
            # Mejorar contraste
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Mejorar nitidez
            image = image.filter(ImageFilter.SHARPEN)
            
            # Guardar imagen preprocesada
            temp_path = image_path.replace('.', '_processed.')
            image.save(temp_path)
            
            return temp_path
        
    except Exception as e:
        print(f"Error en preprocesamiento: {e}")
        return image_path  # Retornar imagen original si hay error

def extract_text_with_ocr(image_path, output_format='word'):
    """
    Extrae texto de una imagen usando OCR con preprocesamiento
    """
    try:
        # Preprocesar imagen
        processed_path = preprocess_image_for_ocr(image_path)
        
        # Configurar Tesseract para mejor precisión
        custom_config = r'--oem 3 --psm 6 -l spa+eng'
        
        # Extraer texto
        extracted_text = pytesseract.image_to_string(processed_path, config=custom_config)
        
        # Limpiar archivo temporal si se creó
        if processed_path != image_path and os.path.exists(processed_path):
            os.remove(processed_path)
        
        return extracted_text.strip()
        
    except Exception as e:
        print(f"Error en OCR: {e}")
        return ""

def create_pdf_from_text(text, image_path, output_path):
    """
    Crea un PDF con el texto extraído y la imagen original
    """
    try:
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        
        # Título
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, "Documento generado por OCR")
        
        # Imagen original
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 100, "Imagen original:")
        
        # Redimensionar imagen para que quepa en la página
        img = ImageReader(image_path)
        img_width, img_height = img.getSize()
        max_width = width - 100
        max_height = 200
        
        # Calcular dimensiones manteniendo proporción
        ratio = min(max_width/img_width, max_height/img_height)
        new_width = img_width * ratio
        new_height = img_height * ratio
        
        c.drawImage(img, 50, height - 100 - new_height - 20, width=new_width, height=new_height)
        
        # Texto extraído
        c.drawString(50, height - 100 - new_height - 50, "Texto extraído:")
        
        # Dividir texto en líneas y agregar al PDF
        lines = text.split('\n')
        y_position = height - 100 - new_height - 80
        
        c.setFont("Helvetica", 10)
        for line in lines:
            if y_position < 50:  # Nueva página si no hay espacio
                c.showPage()
                y_position = height - 50
                c.setFont("Helvetica", 10)
            
            # Dividir líneas largas
            if len(line) > 80:
                words = line.split()
                current_line = ""
                for word in words:
                    if len(current_line + word) > 80:
                        c.drawString(50, y_position, current_line)
                        y_position -= 15
                        current_line = word + " "
                    else:
                        current_line += word + " "
                if current_line:
                    c.drawString(50, y_position, current_line)
                    y_position -= 15
            else:
                c.drawString(50, y_position, line)
                y_position -= 15
        
        c.save()
        return True
        
    except Exception as e:
        print(f"Error creando PDF: {e}")
        return False

@csrf_protect
@require_http_methods(["GET", "POST"])
def convertir_documento(request):
    message = None
    archivo_convertido = None
    
    if request.method == 'POST' and request.FILES.get('document_file'):
        document_file = request.FILES['document_file']
        
        # Validar tamaño del archivo
        if document_file.size > settings.MAX_UPLOAD_SIZE:
            message = f"El archivo es demasiado grande. Máximo permitido: {settings.MAX_UPLOAD_SIZE // (1024*1024)}MB"
            return render(request, 'index.html', {'message': message})
        
        # Validar extensión del archivo
        file_extension = os.path.splitext(document_file.name)[1].lower()
        if file_extension not in settings.ALLOWED_UPLOAD_EXTENSIONS:
            message = f"Tipo de archivo no permitido. Solo se permiten: {', '.join(settings.ALLOWED_UPLOAD_EXTENSIONS)}"
            return render(request, 'index.html', {'message': message})
        
        # Validar tipo MIME según el tipo de archivo
        mime_type, _ = mimetypes.guess_type(document_file.name)
        if file_extension in ['.docx', '.doc']:
            if mime_type not in settings.ALLOWED_WORD_MIMES:
                message = "Tipo de archivo no válido. Solo se permiten documentos de Word válidos."
                return render(request, 'index.html', {'message': message})
        elif file_extension == '.pdf':
            if mime_type not in settings.ALLOWED_PDF_MIMES:
                message = "Tipo de archivo no válido. Solo se permiten archivos PDF válidos."
                return render(request, 'index.html', {'message': message})
        
        # Directorio donde se guardará el archivo
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'converted_files')
        os.makedirs(upload_dir, exist_ok=True)
        fs = FileSystemStorage(location=upload_dir)

        # Sanitizar nombre del archivo
        safe_filename = ''.join(c for c in document_file.name if c.isalnum() or c in '._-')
        if not safe_filename:
            safe_filename = f'documento{file_extension}'
        
        # Guardar el archivo
        filename = fs.save(safe_filename, document_file)
        filepath = os.path.join(upload_dir, filename)

        try:
            if file_extension in ['.docx', '.doc']:
                # Convertir Word a PDF usando LibreOffice
                output_filename = os.path.splitext(filename)[0] + '.pdf'
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', filepath, '--outdir', upload_dir],
                    check=True,
                    timeout=60,  # Timeout de 60 segundos
                    capture_output=True,
                    text=True
                )
                
                # Verificar que el archivo de salida se haya creado
                output_path = os.path.join(upload_dir, output_filename)
                if not os.path.exists(output_path):
                    raise Exception("LibreOffice no pudo generar el archivo PDF")
                
                conversion_type = "Word a PDF"
            
            elif file_extension == '.pdf':
                # Convertir PDF a Word usando pdf2docx
                output_filename = os.path.splitext(filename)[0] + '.docx'
                output_path = os.path.join(upload_dir, output_filename)
                
                cv = Converter(filepath)
                cv.convert(output_path, start=0, end=None)
                cv.close()
                conversion_type = "PDF a Word"
            
            # Conversión de imagen usando OCR mejorado
            elif file_extension in ['.jpg', '.jpeg', '.png']:
                if mime_type not in settings.ALLOWED_IMAGE_MIMES:
                    message = "Tipo de archivo de imagen no válido."
                    return render(request, 'index.html', {'message': message})
                
                try:
                    # Abrir la imagen con PIL y validar formato
                    image = Image.open(filepath)
                    
                    # Verificar que es realmente una imagen válida
                    image.verify()
                    
                    # Reabrir la imagen después de verify() (que la cierra)
                    image = Image.open(filepath)
                    
                    # Validar dimensiones mínimas y máximas
                    width, height = image.size
                    if width < 50 or height < 50:
                        message = "La imagen es demasiado pequeña. Dimensiones mínimas: 50x50 píxeles."
                        return render(request, 'index.html', {'message': message})
                    
                    if width > 4000 or height > 4000:
                        message = "La imagen es demasiado grande. Dimensiones máximas: 4000x4000 píxeles."
                        return render(request, 'index.html', {'message': message})
                    
                    # Convertir a RGB si es necesario (para compatibilidad con OCR)
                    if image.mode not in ('RGB', 'L'):
                        image = image.convert('RGB')
                        
                except Exception as e:
                    message = f"Error al procesar la imagen: {str(e)}. Asegúrese de que el archivo sea una imagen válida."
                    return render(request, 'index.html', {'message': message})
                
                # Obtener formato de salida del formulario
                output_format = request.POST.get('output_format', 'word').lower()
                
                # Extraer texto usando OCR mejorado
                extracted_text = extract_text_with_ocr(filepath, output_format)
                
                if not extracted_text.strip():
                    message = "No se pudo extraer texto de la imagen. Asegúrese de que la imagen contenga texto legible y tenga buena calidad."
                    return render(request, 'index.html', {'message': message})
                
                if output_format == 'pdf':
                    # Crear PDF con texto extraído
                    output_filename = os.path.splitext(filename)[0] + '.pdf'
                    output_path = os.path.join(upload_dir, output_filename)
                    
                    success = create_pdf_from_text(extracted_text, filepath, output_path)
                    if not success:
                        message = "Error al crear el archivo PDF."
                        return render(request, 'index.html', {'message': message})
                    
                    conversion_type = "Imagen a PDF (OCR)"
                    
                else:  # Formato Word por defecto
                    # Crear documento Word
                    output_filename = os.path.splitext(filename)[0] + '.docx'
                    output_path = os.path.join(upload_dir, output_filename)
                    
                    doc = Document()
                    doc.add_heading('Texto extraído de imagen', 0)
                    
                    # Agregar la imagen al documento
                    doc.add_heading('Imagen original:', level=1)
                    doc.add_picture(filepath, width=Inches(6))
                    
                    # Agregar el texto extraído
                    doc.add_heading('Texto extraído:', level=1)
                    
                    # Dividir el texto en párrafos
                    paragraphs = extracted_text.split('\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():  # Solo agregar párrafos no vacíos
                            doc.add_paragraph(paragraph.strip())
                    
                    # Guardar el documento
                    doc.save(output_path)
                    
                    conversion_type = "Imagen a Word (OCR)"
            
            # Eliminar el archivo original después de la conversión
            os.remove(filepath)

            archivo_convertido = output_filename
            message = f"¡Se ha convertido {document_file.name} de {conversion_type} exitosamente!"
            
        except subprocess.TimeoutExpired:
            message = "Error: La conversión tardó demasiado tiempo. Intenta con un archivo más pequeño."
        except subprocess.CalledProcessError as e:
            if 'libreoffice' in str(e.cmd):
                message = "Error: No se pudo convertir el archivo DOCX. Verifica que el archivo no esté corrupto."
            else:
                message = f"Error en la conversión: {e}"
        except Exception as e:
            message = f"Error inesperado en la conversión: {e}"

    return render(request, 'index.html', {'message': message, 'archivo_convertido': archivo_convertido})


@require_http_methods(["GET"])
def descargar_archivo(request, filename):
    """Permite descargar el archivo convertido (PDF o DOCX)."""
    file_path = os.path.join(settings.MEDIA_ROOT, 'converted_files', filename)

    if os.path.exists(file_path):
        # Determinar el tipo de contenido basado en la extensión
        file_extension = os.path.splitext(filename)[1].lower()
        if file_extension == '.pdf':
            content_type = 'application/pdf'
        elif file_extension == '.docx':
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_extension in ['.jpg', '.jpeg']:
            content_type = 'image/jpeg'
        elif file_extension == '.png':
            content_type = 'image/png'
        else:
            content_type = 'application/octet-stream'
        
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
    else:
        return HttpResponseNotFound('El archivo convertido no se encontró')
